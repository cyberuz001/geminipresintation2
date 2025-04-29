import io
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from image_scrapper import downloader
except ImportError:
    from .image_scrapper import downloader


async def generate_docx_prompt(language, emotion_type, topic, page_count="4"):
    message = f"""Create an {language} language very long outline for a {emotion_type} research paper on the topic of {topic} which is long as much as possible and should be approximately {page_count} pages long. 
Language of research paper - {language}.
Provide as much information as possible.

Put this tag before the Title: [TITLE]
Put this tag after the Title: [/TITLE]
Put this tag before the Subtitle: [SUBTITLE]
Put this tag after the Subtitle: [/SUBTITLE]
Put this tag before the Heading: [HEADING]
Put this tag after the Heading: [/HEADING]
Put this tag before the Content: [CONTENT]
Put this tag after the Content: [/CONTENT]
Put this tag before the Image: [IMAGE]
Put this tag after the Image: [/IMAGE]

Elaborate on the Content, provide as much information as possible.
You put a [/CONTENT] at the end of the Content.
Do not put a tag before ending previous.

IMPORTANT STRUCTURE REQUIREMENTS:
1. Create at least {int(page_count) + 1} different headings to ensure enough content for {page_count} pages
2. After EACH heading and its content, add at least one relevant image tag
3. Make sure to distribute content evenly across all {page_count} pages
4. Each heading should have substantial content (at least 200-300 words)

For example:
[TITLE]Mental Health[/TITLE]
[SUBTITLE]Understanding and Nurturing Your Mind: A Guide to Mental Health[/SUBTITLE]
[HEADING]Mental Health Definition[/HEADING]
[CONTENT]...(detailed content here, at least 200-300 words)...[/CONTENT]
[IMAGE]Person Meditating Mental Health[/IMAGE]
[HEADING]Types of Mental Health Disorders[/HEADING]
[CONTENT]...(detailed content here, at least 200-300 words)...[/CONTENT]
[IMAGE]Diagram of Mental Health Disorders[/IMAGE]

Pay attention to the language of research paper - {language}.
IMPORTANT: Always describe images in English regardless of the paper language, as this will yield better image search results.
Each image should be described in detail with specific keywords, such as "Mount Everest Sunset Himalaya Mountains" or "Niagara Falls Rainbow Waterfall".
Do not reply as if you are talking about the research paper itself. (ex. "Include pictures here about...")
Do not include any special characters (?, !, ., :, ) in the Title.
Do not include any additional information in your response and stick to the format.

IMPORTANT: Make sure the content is substantial enough to fill approximately {page_count} pages in a document."""

    return message


async def generate_docx(answer, page_count="4"):
    doc = Document()
    
    # Set document properties for better formatting
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    async def split_tags(reply):
        pattern = r'\[(.*?)\](.*?)\[/\1\]'
        tags = re.findall(pattern, reply, re.DOTALL)
        return tags

    async def parse_response(tags_array, target_pages):
        if not tags_array:
            raise IndexError
        
        # Add title with special formatting
        for item in tags_array:
            if item[0] == 'TITLE':
                title = doc.add_heading(item[1], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
        
        # Add subtitle with special formatting
        for item in tags_array:
            if item[0] == 'SUBTITLE':
                subtitle = doc.add_paragraph()
                subtitle_run = subtitle.add_run(item[1])
                subtitle_run.italic = True
                subtitle_run.font.size = Pt(14)
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
        
        # Extract all headings, contents and images for better organization
        headings = []
        contents = []
        images = []
        
        for item in tags_array:
            if item[0] == 'HEADING':
                headings.append(item[1])
            elif item[0] == 'CONTENT':
                contents.append(item[1])
            elif item[0] == 'IMAGE':
                images.append(item[1])
        
        # Ensure we have enough content for the target number of pages
        # If we don't have enough headings/content, we'll repeat some
        min_sections_needed = max(int(target_pages) * 1, 1)  # At least 1 section per page
        
        # Process the content in sections (heading + content + image)
        section_count = min(len(headings), len(contents))
        
        for i in range(section_count):
            # Add heading
            heading = doc.add_heading(headings[i], 2)
            heading.style.font.size = Pt(14)
            
            # Add content
            paragraphs = contents[i].split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.style.font.size = Pt(12)
            
            # Add image if available
            img_index = min(i, len(images) - 1) if images else -1
            if img_index >= 0:
                try:
                    # Always search in English for better results
                    image_query = images[img_index]
                    # Add some general keywords to improve search results
                    enhanced_query = f"{image_query} high quality photo"
                    
                    image_data = await downloader.download(
                        enhanced_query, 
                        limit=1, 
                        adult_filter_off=True, 
                        timeout=15,
                        filter="+filterui:aspect-wide+filterui:imagesize-wallpaper+filterui:photo-photo"
                    )
                    
                    # Add image with proper sizing
                    doc.add_picture(io.BytesIO(image_data), width=Inches(6))
                    
                    # Add image caption
                    caption = doc.add_paragraph(f"Figure: {image_query}")
                    caption.style.font.italic = True
                    caption.style.font.size = Pt(10)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add some space after the image
                    doc.add_paragraph()
                except Exception as e:
                    print(f"Error downloading image: {e}")
                    # Add a placeholder paragraph if image fails
                    doc.add_paragraph("Image could not be loaded.")

    async def find_title(tags_array):
        for item in tags_array:
            if item[0] == 'TITLE':
                return item[1]
        return "Abstract"

    reply_array = await split_tags(answer)
    await parse_response(reply_array, int(page_count))
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_title = f"{await find_title(reply_array)}.docx"
    print(f"done {docx_title}")

    return docx_bytes, docx_title
